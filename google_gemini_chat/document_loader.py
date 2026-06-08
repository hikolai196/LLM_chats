from io import BytesIO, StringIO
from pathlib import Path

import pandas as pd
from docx import Document


SUPPORTED_EXTENSIONS = {".csv", ".docx", ".txt", ".xls", ".xlsx"}


def load_document(file_name, file_bytes):
    """Extract readable text from an uploaded document."""
    extension = Path(file_name).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}")

    if extension == ".docx":
        return _load_docx(file_bytes)

    if extension in {".xls", ".xlsx"}:
        return _load_excel(file_bytes)

    if extension == ".csv":
        return _load_csv(file_bytes)

    return _load_txt(file_bytes)


def build_document_context(documents):
    """Combine loaded documents into a context block for the LLM."""
    context_parts = []

    for document in documents:
        context_parts.append(
            f"### Document: {document['name']}\n{document['content']}"
        )

    return "\n\n".join(context_parts)


def _decode_text(file_bytes):
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    return file_bytes.decode("utf-8", errors="replace")


def _load_txt(file_bytes):
    return _decode_text(file_bytes)


def _load_csv(file_bytes):
    text = _decode_text(file_bytes)
    dataframe = pd.read_csv(StringIO(text))
    return dataframe.to_csv(index=False)


def _load_docx(file_bytes):
    document = Document(BytesIO(file_bytes))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_text = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(cells))

    return "\n".join(paragraphs + table_text)


def _load_excel(file_bytes):
    sheets = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
    sheet_text = []

    for sheet_name, dataframe in sheets.items():
        sheet_text.append(f"## Sheet: {sheet_name}")
        sheet_text.append(dataframe.to_csv(index=False))

    return "\n\n".join(sheet_text)
